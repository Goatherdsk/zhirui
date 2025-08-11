#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
智睿商务车产品数据转换脚本
将DOCX文件转换为JSON格式的产品数据
"""

import os
import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches
import base64
from PIL import Image
import io
import logging

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class ProductConverter:
    def __init__(self, input_dir='public/products', output_dir='public/data'):
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        self.images_dir = Path('public/images/products')
        
        # 确保输出目录存在
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.images_dir.mkdir(parents=True, exist_ok=True)
        
        # 产品数据结构模板
        self.product_template = {
            "id": "",
            "name": "",
            "subtitle": "",
            "category": "",
            "series": "",
            "exteriorColor": "",
            "interiorColor": "",
            "price": "",
            "originalPrice": "",
            "description": "",
            "features": [],
            "specs": {
                "engine": "",
                "power": "",
                "transmission": "",
                "fuel": "",
                "seats": "",
                "length": "",
                "width": "",
                "height": "",
                "wheelbase": "",
                "displacement": "",
                "maxSpeed": "",
                "acceleration": ""
            },
            "highlights": [],
            "images": {
                "exterior": [],
                "interior": [],
                "details": []
            },
            "gradient": "",
            "badge": "",
            "availability": "available",
            "launchDate": "",
            "warranty": ""
        }
        
        # 颜色映射
        self.color_gradients = {
            "曜影黑": "linear-gradient(135deg, #1a1a1a 0%, #2d2d2d 100%)",
            "幻影白": "linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%)",
            "尊行版": "linear-gradient(135deg, #c9a96e 0%, #f4e4bc 100%)",
            "凡戴克咖": "linear-gradient(135deg, #8b4513 0%, #a0522d 100%)",
            "凡戴克棕": "linear-gradient(135deg, #654321 0%, #8b4513 100%)",
            "勃艮第红": "linear-gradient(135deg, #800020 0%, #b03060 100%)",
            "帕米尔黄": "linear-gradient(135deg, #ffd700 0%, #ffeb3b 100%)",
            "爱马仕橙": "linear-gradient(135deg, #ff6b35 0%, #ff8c42 100%)",
            "牡蛎白": "linear-gradient(135deg, #faf0e6 0%, #f5f5dc 100%)",
            "琥珀棕": "linear-gradient(135deg, #cd853f 0%, #daa520 100%)"
        }
        
        # 车型分类映射
        self.category_mapping = {
            "尊行版": "premium",
            "行政版": "executive", 
            "科技版": "tech",
            "至尊版": "luxury"
        }

    def extract_images_from_docx(self, doc_path, product_id):
        """从DOCX文件中提取图片"""
        doc = Document(doc_path)
        images = {"exterior": [], "interior": [], "details": []}
        image_counter = 0
        
        try:
            # 遍历文档中的所有关系
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        image_counter += 1
                        
                        # 生成图片文件名
                        image_filename = f"{product_id}_image_{image_counter}.jpg"
                        image_path = self.images_dir / image_filename
                        
                        # 保存图片
                        with open(image_path, 'wb') as img_file:
                            img_file.write(image_data)
                        
                        # 根据图片位置和内容判断类型
                        relative_path = f"/images/products/{image_filename}"
                        if image_counter <= 3:
                            images["exterior"].append(relative_path)
                        elif image_counter <= 6:
                            images["interior"].append(relative_path)
                        else:
                            images["details"].append(relative_path)
                            
                        logger.info(f"提取图片: {image_filename}")
                        
                    except Exception as e:
                        logger.warning(f"提取图片失败: {e}")
                        
        except Exception as e:
            logger.error(f"处理图片时出错: {e}")
            
        return images

    def parse_filename(self, filename):
        """解析文件名获取产品信息"""
        # 移除扩展名
        name_without_ext = filename.replace('.docx', '')
        
        # 解析文件名 例如: "尊行版 曜影黑，凡戴克咖"
        parts = name_without_ext.split('，')
        
        if len(parts) >= 2:
            # 外观色和内饰色
            first_part = parts[0].strip()
            interior_color = parts[1].strip()
            
            # 分离系列和外观色
            series_color_parts = first_part.split(' ')
            if len(series_color_parts) >= 2:
                series = series_color_parts[0]
                exterior_color = ' '.join(series_color_parts[1:])
            else:
                series = "标准版"
                exterior_color = first_part
                
        else:
            # 只有外观色的情况
            color_parts = name_without_ext.split(' ')
            if len(color_parts) >= 2:
                series = color_parts[0]
                exterior_color = ' '.join(color_parts[1:])
                interior_color = "标准内饰"
            else:
                series = "标准版"
                exterior_color = name_without_ext
                interior_color = "标准内饰"
        
        return series, exterior_color, interior_color

    def extract_text_content(self, doc_path):
        """提取DOCX文档中的文本内容"""
        try:
            doc = Document(doc_path)
            content = {
                "title": "",
                "description": "",
                "features": [],
                "specs": {},
                "price": "",
                "highlights": []
            }
            
            all_text = []
            
            # 提取所有段落文本
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    all_text.append(text)
            
            # 提取表格数据
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        all_text.append(" | ".join(row_text))
            
            # 解析内容
            content = self.parse_content(all_text)
            
            return content
            
        except Exception as e:
            logger.error(f"提取文本内容失败: {e}")
            return {
                "title": "",
                "description": "",
                "features": [],
                "specs": {},
                "price": "",
                "highlights": []
            }

    def parse_content(self, text_lines):
        """解析文档内容"""
        content = {
            "title": "",
            "description": "",
            "features": [],
            "specs": {},
            "price": "",
            "highlights": []
        }
        
        # 合并所有文本
        full_text = "\n".join(text_lines)
        
        # 提取价格信息
        price_patterns = [
            r'价格[：:]\s*(\d+\.?\d*万?)',
            r'售价[：:]\s*(\d+\.?\d*万?)',
            r'(\d+\.?\d*万)起',
            r'￥\s*(\d+\.?\d*万?)'
        ]
        
        for pattern in price_patterns:
            match = re.search(pattern, full_text)
            if match:
                content["price"] = match.group(1) + "起"
                break
        
        # 提取规格参数
        spec_patterns = {
            "engine": [r'发动机[：:](.+?)(?:\n|$)', r'引擎[：:](.+?)(?:\n|$)'],
            "power": [r'功率[：:](.+?)(?:\n|$)', r'马力[：:](.+?)(?:\n|$)'],
            "transmission": [r'变速箱[：:](.+?)(?:\n|$)', r'变速器[：:](.+?)(?:\n|$)'],
            "fuel": [r'油耗[：:](.+?)(?:\n|$)', r'百公里[：:](.+?)(?:\n|$)'],
            "seats": [r'座位[：:](.+?)(?:\n|$)', r'座椅[：:](.+?)(?:\n|$)'],
            "length": [r'长度[：:](.+?)(?:\n|$)', r'车长[：:](.+?)(?:\n|$)'],
            "width": [r'宽度[：:](.+?)(?:\n|$)', r'车宽[：:](.+?)(?:\n|$)'],
            "height": [r'高度[：:](.+?)(?:\n|$)', r'车高[：:](.+?)(?:\n|$)'],
            "wheelbase": [r'轴距[：:](.+?)(?:\n|$)'],
            "displacement": [r'排量[：:](.+?)(?:\n|$)'],
            "maxSpeed": [r'最高速度[：:](.+?)(?:\n|$)', r'极速[：:](.+?)(?:\n|$)'],
            "acceleration": [r'加速[：:](.+?)(?:\n|$)', r'百公里加速[：:](.+?)(?:\n|$)']
        }
        
        for key, patterns in spec_patterns.items():
            for pattern in patterns:
                match = re.search(pattern, full_text)
                if match:
                    content["specs"][key] = match.group(1).strip()
                    break
        
        # 提取配置特性
        feature_keywords = [
            "真皮", "按摩", "加热", "通风", "记忆", "电动", "自动", "智能", "导航", 
            "音响", "空调", "天窗", "LED", "倒车", "雷达", "摄像", "蓝牙", "USB",
            "无线充电", "空气净化", "氛围灯", "隐私玻璃", "商务", "办公"
        ]
        
        for line in text_lines:
            for keyword in feature_keywords:
                if keyword in line and len(line) < 50:  # 避免太长的描述
                    if line not in content["features"]:
                        content["features"].append(line.strip())
        
        # 生成描述
        if not content["description"]:
            content["description"] = "高端商务车型，集舒适、豪华、科技于一体，为您提供卓越的出行体验。"
        
        # 生成亮点
        if content["price"]:
            content["highlights"].append("性价比之选")
        if any("智能" in feature for feature in content["features"]):
            content["highlights"].append("智能科技")
        if any("真皮" in feature for feature in content["features"]):
            content["highlights"].append("豪华内饰")
        
        return content

    def generate_product_id(self, series, exterior_color, interior_color):
        """生成产品ID"""
        # 创建简短的ID
        series_short = {
            "尊行版": "zx",
            "行政版": "xz", 
            "科技版": "kj",
            "至尊版": "zz"
        }.get(series, "std")
        
        exterior_short = {
            "曜影黑": "yh",
            "幻影白": "hw",
            "勃艮第红": "bgd",
            "帕米尔黄": "pme",
            "爱马仕橙": "ams",
            "牡蛎白": "mb",
            "琥珀棕": "hp"
        }.get(exterior_color, "def")
        
        interior_short = {
            "凡戴克咖": "fdk",
            "凡戴克棕": "fdz",
            "勃艮第红": "bgd", 
            "帕米尔黄": "pme",
            "爱马仕橙": "ams",
            "牡蛎白": "mb",
            "琥珀棕": "hp"
        }.get(interior_color, "std")
        
        return f"{series_short}_{exterior_short}_{interior_short}"

    def convert_single_docx(self, docx_path):
        """转换单个DOCX文件"""
        filename = docx_path.name
        logger.info(f"处理文件: {filename}")
        
        try:
            # 解析文件名
            series, exterior_color, interior_color = self.parse_filename(filename)
            
            # 生成产品ID
            product_id = self.generate_product_id(series, exterior_color, interior_color)
            
            # 提取文档内容
            content = self.extract_text_content(docx_path)
            
            # 提取图片
            images = self.extract_images_from_docx(docx_path, product_id)
            
            # 构建产品数据
            product = self.product_template.copy()
            product.update({
                "id": product_id,
                "name": f"智锐·{series}",
                "subtitle": f"{exterior_color} × {interior_color}",
                "category": self.category_mapping.get(series, "executive"),
                "series": series,
                "exteriorColor": exterior_color,
                "interiorColor": interior_color,
                "price": content["price"] or "面议",
                "description": content["description"],
                "features": content["features"][:8],  # 限制特性数量
                "specs": content["specs"],
                "highlights": content["highlights"],
                "images": images,
                "gradient": self.color_gradients.get(exterior_color, 
                          self.color_gradients.get(series, 
                          "linear-gradient(135deg, #2d2d2d 0%, #1a1a1a 100%)")),
                "badge": self.get_badge(series),
                "launchDate": "2024-01-01",
                "warranty": "3年或10万公里"
            })
            
            return product
            
        except Exception as e:
            logger.error(f"转换文件 {filename} 失败: {e}")
            return None

    def get_badge(self, series):
        """根据系列获取徽章"""
        badge_mapping = {
            "尊行版": "PREMIUM",
            "行政版": "EXECUTIVE",
            "科技版": "TECH",
            "至尊版": "LUXURY"
        }
        return badge_mapping.get(series, "NEW")

    def convert_all_products(self):
        """转换所有产品文件"""
        logger.info("开始转换产品数据...")
        
        products = []
        docx_files = list(self.input_dir.glob("*.docx"))
        
        if not docx_files:
            logger.warning("未找到DOCX文件")
            return
        
        logger.info(f"找到 {len(docx_files)} 个产品文件")
        
        for docx_path in docx_files:
            product = self.convert_single_docx(docx_path)
            if product:
                products.append(product)
                logger.info(f"成功转换: {product['name']} - {product['subtitle']}")
        
        # 保存产品数据
        output_file = self.output_dir / "products.json"
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(products, f, ensure_ascii=False, indent=2)
        
        logger.info(f"转换完成! 共处理 {len(products)} 个产品")
        logger.info(f"数据已保存到: {output_file}")
        
        # 生成产品摘要
        self.generate_summary(products)
        
        return products

    def generate_summary(self, products):
        """生成产品摘要"""
        summary = {
            "totalProducts": len(products),
            "categories": {},
            "series": {},
            "colors": {
                "exterior": {},
                "interior": {}
            },
            "priceRange": {
                "min": None,
                "max": None
            },
            "lastUpdated": "2024-01-01"
        }
        
        # 统计分类
        for product in products:
            category = product["category"]
            series = product["series"]
            exterior = product["exteriorColor"]
            interior = product["interiorColor"]
            
            summary["categories"][category] = summary["categories"].get(category, 0) + 1
            summary["series"][series] = summary["series"].get(series, 0) + 1
            summary["colors"]["exterior"][exterior] = summary["colors"]["exterior"].get(exterior, 0) + 1
            summary["colors"]["interior"][interior] = summary["colors"]["interior"].get(interior, 0) + 1
        
        # 保存摘要
        summary_file = self.output_dir / "products_summary.json"
        with open(summary_file, 'w', encoding='utf-8') as f:
            json.dump(summary, f, ensure_ascii=False, indent=2)
        
        logger.info(f"产品摘要已保存到: {summary_file}")

def main():
    """主函数"""
    print("🚗 智睿商务车产品数据转换工具")
    print("=" * 50)
    
    converter = ProductConverter()
    
    try:
        products = converter.convert_all_products()
        
        print(f"\n✅ 转换完成!")
        print(f"📊 共处理 {len(products)} 个产品")
        print(f"📁 数据保存在: public/data/")
        print(f"🖼️  图片保存在: public/images/products/")
        
    except Exception as e:
        print(f"\n❌ 转换失败: {e}")
        logger.error(f"转换过程中出现错误: {e}")

if __name__ == "__main__":
    main()
